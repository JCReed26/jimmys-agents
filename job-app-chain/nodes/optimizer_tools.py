"""
Drive I/O and DOCX editing tools for the optimizer agent.
"""

import os
import tempfile
from pathlib import Path

from langchain_core.tools import tool
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from docx import Document
import io

SCOPES = [
    "https://www.googleapis.com/auth/drive",
]


def _get_drive_service():
    service_account_file = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
    if service_account_file and os.path.exists(service_account_file):
        creds = Credentials.from_service_account_file(service_account_file, scopes=SCOPES)
    else:
        token_paths = ["secrets/sheets_token.json", "secrets/token.json", "token.json"]
        from google.oauth2.credentials import Credentials as UserCredentials
        creds = None
        for path in token_paths:
            if os.path.exists(path):
                creds = UserCredentials.from_authorized_user_file(path, scopes=SCOPES)
                break
        if not creds:
            raise RuntimeError("No valid Google credentials found for Drive access")
    return build("drive", "v3", credentials=creds)


@tool
def download_drive_template(file_id: str) -> str:
    """Download a Google Doc as .docx and return the local file path."""
    service = _get_drive_service()
    request = service.files().export_media(
        fileId=file_id,
        mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    downloader = MediaIoBaseDownload(tmp, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    tmp.close()
    return tmp.name


@tool
def upload_to_drive(local_path: str, filename: str, folder_id: str) -> str:
    """Upload a local file to a Google Drive folder and return a shareable URL."""
    service = _get_drive_service()
    file_metadata = {"name": filename, "parents": [folder_id]}
    media = MediaFileUpload(
        local_path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    uploaded = service.files().create(
        body=file_metadata,
        media_body=media,
        fields="id"
    ).execute()
    file_id = uploaded.get("id")
    service.permissions().create(
        fileId=file_id,
        body={"type": "anyone", "role": "reader"}
    ).execute()
    return f"https://drive.google.com/file/d/{file_id}/view"


@tool
def fill_docx_template(template_path: str, replacements: dict) -> str:
    """Replace {{PLACEHOLDER}} variables in a .docx file and return the output path.

    Args:
        template_path: Path to the source .docx template file.
        replacements: Dict mapping placeholder names (without braces) to replacement text.
            e.g. {"CANDIDATE_NAME": "Jimmy", "JOB_TITLE": "Software Engineer"}

    Returns:
        Path to the filled output .docx file.
    """
    doc = Document(template_path)

    def _replace_in_run(run):
        for key, value in replacements.items():
            token = "{{" + key + "}}"
            if token in run.text:
                run.text = run.text.replace(token, str(value))

    for para in doc.paragraphs:
        for run in para.runs:
            _replace_in_run(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _replace_in_run(run)

    output_path = template_path.replace(".docx", "_filled.docx")
    doc.save(output_path)
    return output_path
